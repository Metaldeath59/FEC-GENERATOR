# Librairies
import pandas as pd
import re
from datetime import datetime
import unicodedata
import numpy as np
from docx import Document
from docx.shared import RGBColor
import os
import locale
from decimal import Decimal
from sys import argv
from pathlib import WindowsPath

locale.setlocale(locale.LC_NUMERIC, 'fr_FR')

# VIABLES SYS INITIALES

df_colonnes_var_1 = [
    "JournalCode",
    "JournalLib",
    "EcritureNum",
    "EcritureDate",
    "CompteNum",
    "CompteLib",
    "CompAuxNum",
    "CompAuxLib",
    "PieceRef",
    "PieceDate",
    "EcritureLib",
    "Debit",
    "Credit",
    "EcritureLet",
    "DateLet",
    "ValidDate",
    "Montantdevise",
    "Idevise"
]

df_colonnes_var_2 = [
    "JournalCode",
    "JournalLib",
    "EcritureNum",
    "EcritureDate",
    "CompteNum",
    "CompteLib",
    "CompAuxNum",
    "CompAuxLib",
    "PieceRef",
    "PieceDate",
    "EcritureLib",
    "Montant",
    "Sens",
    "EcritureLet",
    "DateLet",
    "ValidDate",
    "Montantdevise",
    "Idevise"
]

variables_success = []
variables_failure = []


# FUNCTIONS

# 2-	Vérification de l’absence de blancs dans les champs obligatoires – Emettre un warning si c’est le cas

def elements_non_autorises_present(liste, elements_autorises):
    # Utilisation d'une liste en compréhension pour obtenir les éléments non autorisés
    elements_non_autorises = [element for element in liste if element not in elements_autorises]

    # Si la liste des éléments non autorisés n'est pas vide, retourner True, sinon False
    return bool(elements_non_autorises)


def verifier_colonnes_vides(df2):
    df = df2.copy()
    colonnes_option = [
        "EcritureLet",
        "DateLet",
        "Montantdevise",
        "Idevise",
        "CompAuxNum",
        "CompAuxLib",
    ]
    try:
        colonnes_vides = df.columns[df.isnull().any()].tolist()
        if colonnes_vides:
            def_2_erreur = f"Erreur: Les colonnes suivantes contiennent des valeurs vides: {', '.join(colonnes_vides)}\
            .\n Les champs autres que {colonnes_option} doivent impérativement être complets, il convient que vous les complétiez."
            return def_2_erreur
        else:
            if elements_non_autorises_present(colonnes_vides, colonnes_option) == False or colonnes_vides == []:
                def_2_ok = "Les champs obligatoires ne contiennent pas de valeurs vides"
            return def_2_ok

    except Exception as e:
        message_erreur = f"Erreur inattendue: {e}"
        return message_erreur
    finally:
        pass


def check_date_format(date_str):
    try:
        datetime.strptime(date_str, '%Y%m%d')
        return True
    except ValueError:
        return False


# 3-	Identification des écritures postérieures à la date limite de remise de la liasse lorsque le champ ValidDate 
# est exploitable– Emettre un warning si c’est le cas

def verifier_dates(df2, liasse):
    df = df2.copy()
    df['FormatCorrect'] = df['EcritureDate'].astype(str).apply(check_date_format)
    if not df['FormatCorrect'].all():
        def_3_erreur_1 = "Erreur format: Certaines dates de la colonne EcritureDate ne sont pas au format YYYYMMDD, vous devez impérativement chager les valeurs du champ EcritureDate"
    else:
        def_3_erreur_1 = "Les formats de la colonne EcritureDate sont corrects"

    df['FormatCorrect'] = df['PieceDate'].astype(str).apply(check_date_format)
    if not df['FormatCorrect'].all():
        def_3_erreur_2 = "Erreur format: Certaines dates de la colonne PieceDate ne sont pas au format YYYYMMDD vous devez impérativement chager les valeurs du champ PieceDate"
    else:
        def_3_erreur_2 = "Les formats de la colonne PieceDate sont corrects"

    df['FormatCorrect'] = df['ValidDate'].astype(str).apply(check_date_format)
    if not df['FormatCorrect'].all():
        def_3_erreur_3 = "Erreur format: Certaines dates de la colonne ValidDate ne sont pas au format YYYYMMDD vous devez impérativement chager les valeurs du champ ValidDate"
    else:
        def_3_erreur_3 = "Les formats de la colonne ValidDate sont corrects"

    liasse_date = pd.to_datetime(liasse, format='%Y%m%d')

    if (pd.to_numeric(df['ValidDate'], errors='coerce') > int(liasse)).any():
        def_3_erreur_4 = "Erreur de date: Certaines dates de ValidDate sont postérieures à la date limite de remise de la liasse. Assurez-vous de pouvoir justifier la saisie tardive de ces montants en cas de contrôle."
    else:
        def_3_erreur_4 = "Aucune des dates de ValidDate ne sont postérieures à la date limite"

    return def_3_erreur_1, def_3_erreur_2, def_3_erreur_3, def_3_erreur_4


# 4-	Vérifier la relation 1 -1 du champ CompteNum et du champ CompteLib – Emettre un warning si c’est le cas

def verifier_relation_1_1(df2):
    unique_val_count = df2.copy().groupby('CompteNum')['CompteLib'].nunique()
    if (unique_val_count != 1).any():
        def_4_erreur = "Erreur de relation 1-1: Il existe plus d'une valeur dans le champ CompteLib pour au moins une valeur du champ CompteNum"
        return def_4_erreur
    else:
        def_4_ok = "Il y a exactement une valeur dans le champ CompteLib pour chaque valeur du champ CompteNum"
        return def_4_ok


# 5-	Vérification de la présence de doublons dans un tableau « EcritureLib par EcritureNum » – Emettre un warning si c’est le cas*
def verifier_relation_1_1_ecriture(df2):
    unique_val_count = df2.copy().groupby('EcritureNum')['EcritureLib'].nunique()
    if (unique_val_count != 1).any():
        def_5_erreur = "Erreur de relation 1-1: Il existe plus d'une valeur dans le champ EcritureLib pour au moins une valeur du champ EcritureNum"
        return def_5_erreur
    else:
        def_5_ok = "Il y a exactement une valeur dans le champ EcritureLib pour chaque valeur du champ EcritureNum"
        return def_5_ok


# 6-	Vérification de l’absence de champs uniquement numériques dans le champ EcritureLib– Emettre un warning si c’est le cas
def verifier_ecriture_lib_numerique(df2):
    df = df2.copy()
    if df['EcritureLib'].astype(str).str.isnumeric().any():
        def_6_error = "Erreur: Ma colonne 'EcritureLib' contient au moins une valeur uniquement numérique. Le champ EcritureLib servant à expliciter une opération, celui-ci doit être explicite. Assurez-vous de modifier les champs concernés en cas de contrôle."
        return def_6_error
    else:
        def_6_ok = "Aucune valeur de la colonne 'EcritureLib n'est uniquement numérique, assurez-vous que les libellés sont tous assez explicites pour pouvoir expliquer une opération."
        return def_6_ok


# 7-	Vérification de la séquentialité par journal du champ EcritureNum– Emettre un warning si c’est le cas

# def extraire_chiffres(chaine):
# chiffres = re.sub(r'\D', '', chaine)
# return int(chiffres) if chiffres else None

def extraire_chiffres(chaine):
    if chaine is None:
        return None  # or handle it accordingly
    elif not isinstance(chaine, str):
        chaine = str(chaine)  # Convert to string if not already
    chiffres = re.sub(r'\D', '', chaine)
    return int(chiffres) if chiffres else None


def check_values(df2):
    if len(df2) < 2:
        return ["Le DataFrame doit avoir au moins deux lignes pour effectuer la vérification."]

    df = df2.copy().sort_values(['JournalCode', 'EcritureNum'])

    def_7_errors = []
    def_7_count_errors = 0

    for i in range(1, len(df)):
        current_index = df.index[i]
        previous_index = df.index[i - 1]

        if df.loc[current_index, 'JournalCode'] == df.loc[previous_index, 'JournalCode']:
            current_ecriture_num_str = df.loc[current_index, 'EcritureNum']
            previous_ecriture_num_str = df.loc[previous_index, 'EcritureNum']

            if current_ecriture_num_str != '' and previous_ecriture_num_str != '':
                current_ecriture_num = extraire_chiffres(current_ecriture_num_str)
                previous_ecriture_num = extraire_chiffres(previous_ecriture_num_str)

                if not (
                        current_ecriture_num == previous_ecriture_num or current_ecriture_num == previous_ecriture_num + 1):
                    def_7_errors.append(
                        f"Ligne {current_index}: l'EcritureNum {df.loc[current_index, 'EcritureNum']} du journal {df.loc[current_index, 'JournalCode']} n'est pas séquencée correctement.")
                    def_7_count_errors += 1
    return f"Nombre d'erreurs dans la séquentialité d'EcritureNum par JournalCode : {def_7_count_errors}", def_7_errors


# 8-	Vérification de la séquentialité PieceRef– Emettre un warning si c’est le cas
def verifier_sequentialite_piece_ref(df2):
    if len(df2) < 2:
        return ["Le DataFrame doit avoir au moins deux lignes pour effectuer la vérification."]

    df = df2.copy().sort_values(['JournalCode', 'EcritureNum', 'PieceRef']).replace("", 0)

    def_8_errors = []
    def_8_count_errors = 0

    for i in range(1, len(df)):
        current_index = df.index[i]
        previous_index = df.index[i - 1]

        if df.loc[current_index, 'JournalCode'] == df.loc[previous_index, 'JournalCode'] and df.loc[
            current_index, 'EcritureNum'] == df.loc[previous_index, 'EcritureNum']:
            current_piece_ref = extraire_chiffres(df.loc[current_index, 'PieceRef'])
            previous_piece_ref = extraire_chiffres(df.loc[previous_index, 'PieceRef'])

            if not (current_piece_ref == previous_piece_ref or current_piece_ref == previous_piece_ref + 1):
                def_8_errors.append(
                    f"La référence de pièce {df.loc[current_index, 'PieceRef']} dans l'écriture {df.loc[current_index, 'EcritureNum']} du journal {df.loc[current_index, 'JournalCode']} n'est pas séquencée correctement.")
                def_8_count_errors += 1

    return f"Nombre d'erreurs dans la séquentialité de Piecered par EcritureNum par JournalCode : {def_8_count_errors}", def_8_errors


# 9-	Checking de l’équilibre des champs Debit-Credit ou Montant-Sens

def verifier_colonne_sens(df):
    # Vérification des valeurs dans la colonne 'Sens'
    valid_values = {'D', 'C'}
    invalid_sens = df[~df['Sens'].isin(valid_values)]

    # Retourner True si la colonne 'Sens' ne contient que des valeurs valides, sinon False
    return invalid_sens.empty


def verifier_compensation_montants(df2, format_montants):
    df = df2.copy()
    df3 = df.copy()

    if format_montants == '1':
        # df['Debit'] = df['Debit'].str.replace('-1', 'C')
        # df['Credit'] = df['Credit'].str.replace('1', 'D')
        # df['Debit'] = df['Debit'].astype(str).str.replace(',', '.').astype(float)
        # df['Credit'] = df['Credit'].astype(str).str.replace(',', '.').astype(float)

        debit = sum(map(Decimal, df['Debit'].astype(str).replace(",", '.')))
        credit = sum(map(Decimal, df['Credit'].astype(str).replace(",", '.')))
        difference = debit - credit
        if difference != 0:
            def_9_errors = f"Erreur : les montants passés au débit et au crédit ne se compensent pas (différence de {abs(difference)}€)! Vérifiez l'égalité de vos champs Debit et Credit et préparez-vous à être en mesure de justifier ces écarts, souvent les écarts Débit-Crédit peut être expliqués par un manque d'arrondi à la décimale."
            return (None, def_9_errors)
        else:
            def_9_ok = "Les montants des colonnes Debit et Credit se compensent."
            return (def_9_ok, None)
    else:
        df['Montant'] *= df['Sens'].map({'C': -1, 'D': 1})
        somme_montant = sum(map(Decimal, df['Montant']))
        if somme_montant != 0:
            def_9_errors = f"Erreur : les montants passés au débit et au crédit ne se compensent pas (différence de {abs(somme_montant)}€)! Vérifiez l'égalité de vos champs Debit et Credit et préparez-vous à être en mesure de justifier ces écarts, souvent les écarts Débit-Crédit peut être expliqués par un manque d'arrondi à la deuxième décimale."
            return (None, def_9_errors)
        else:
            def_9_ok = "Les montants de la colonne Montant se compensent entre eux."
            return (def_9_ok, None)


# 10-Lister les numéros d’écritures pour lesquels le montant du compte provision est le plus important– Emettre un warning si c’est le cas
def selectionner_top30_comptes(df2, format_montants):
    messages = []  # Initialisez la liste des messages

    df = df2[df2['CompteNum'].astype(str).str.startswith('15')].copy()

    if format_montants == '1':
        df2 = df2.sort_values(by='Credit', ascending=False).head(15)
        for _, row in df2.iterrows():
            message = f"Il convient de vérifier le montant {row['Credit']} de la PieceRef {row['PieceRef']}."
            messages.append(message)
    else:
        # Vérification de la validité des valeurs dans la colonne 'Sens'
        valid_sens_values = {'C', 'D', 1, -1}
        invalid_sens_mask = ~df2['Sens'].isin(valid_sens_values)

        # Traitement des valeurs invalides ou suppression des lignes
        if invalid_sens_mask.any():
            # Vous pouvez traiter ces valeurs d'une manière spécifique ou simplement les ignorer
            df2.loc[invalid_sens_mask, 'Sens'] = np.nan  # Remplacez les valeurs invalides par NaN

        # Ajout d'une condition pour éviter la multiplication si 'Sens' est déjà 1 ou -1
        df2.loc[~df2['Sens'].isin({1, -1}), 'Montant'] *= df2['Sens'].map({'C': -1, 'D': 1})

        df2 = df2.sort_values(by='Montant', ascending=False).head(15)

        # Supprimer les lignes avec des doublons dans la colonne 'EcritureNum'
        df2 = df2.drop_duplicates(subset='EcritureNum', keep='first')

        for _, row in df2.iterrows():
            message = f"Il convient de vérifier le montant {row['Montant']} de la PieceRef {row['PieceRef']}."
            messages.append(message)

    return messages


# 11-Emettre un warning lorsque le libellé comporte les chaines de caractères interdits

def remove_accents(input_str):
    if isinstance(input_str, str):
        nfkd_form = unicodedata.normalize('NFKD', input_str)
        return ''.join([c for c in nfkd_form if not unicodedata.combining(c)])
    else:
        return input_str


def verifier_valeurs_interdites_ecriture_lib(df2):
    df = df2.copy()
    forbidden_list = [
        'samsung',
        'apple',
        'iphone',
        'telephonie',
        'smartphone',
        'restaurant',
        'repas',
        'diner',
        'dejeuner',
        'essence',
        'voiture',
        'volkswagen',
        'vacances',
        'cadeau',
        'tesla',
        'louisvuitton',
        'rolex',
        'ferrari',
        'cartier',
        'chanel',
        'yacht',
        'helicoptere',
        'caviar',
        'champagne',
        'versace',
        'croisiere',
        'voyage',
        'maserati',
        'cristal',
        'diamants',
        'diner',
        'vip',
        'cognac',
        'golf',
        'limousine',
        'soie',
        'art',
        'platine',
        'michelin',
        'platine',
        'piano',
        'chef',
        'caleche',
        'safari',
        'chalet',
        'ski',
        'luxe',
        'voilier',
        'cuisine',
        'cinema',
        'elite',
        'suite',
        'whisky',
        'avion',
        'millesime',
        'exclusif',
        'collier',
        'helicoptere',
        'bracelet',
        'evenement',
        'croisiere',
        'meuble',
        'conciergerie',
        'location',
        'safari',
        'premium',
        'hotel',
        'escapade',
    ]

    df['EcritureLib'] = df['EcritureLib'].astype(str).apply(
        lambda x: remove_accents(x) if pd.notnull(x) else x).str.lower()

    df['EcritureNum'] = df['EcritureNum'].str.lower()

    detected_values = []

    for forbidden_value in forbidden_list:
        if df['EcritureLib'].str.contains(forbidden_value).any() or df['EcritureNum'].str.contains(
                forbidden_value).any():
            detected_values.append(forbidden_value)

    if detected_values:
        def_11_errors = f"Erreur: La colonne EcritureLib contient les valeurs suspectes suivantes : {', '.join(detected_values)}. Soyez certains de pouvoir justifier les écriutres concernées en fonction de l'activité de votre société. Notez que les mots courts comme 'Art' détectés peuvent être des parties de mots, ces occurences sont à ignorer."
        return def_11_errors
    else:
        def_11_ok = "La colonne EcritureLib ne contient pas de valeur suspecte."
        return def_11_ok


# QUESTIONNAIRE

print("Partie 1: Renseignement des informations relatives à l'entité et l'exercice:\n")

siren = input("Veuillez entrer le SIREN de l'entité concernée:\n")

periode_end = input("Veuillez entrer la date de clôture de l'exercice concerné au format AAAAMMJJ:\n")

liasse = input("Veuillez entrer la date limite du dépôt de la liasse au format AAAAMMJJ: \n")

while True:
    format_fichier = input(
        "Merci de renseigner le format du fichier à convertir:\n1-Excel: Tapez xlsx\n2-Fichier plat: Tapez selon le format csv ou txt:\n")
    if format_fichier in ('xlsx', 'csv', 'txt'):
        break
    else:
        "Merci de vérifier avoir un fichier d'extraction au format valide et de bien le renseigner ici"

if not (format_fichier == 'xlsx'):
    encoding_fichier = input("Merci de renseigner le format d'encodage du fichier plat entré (par exemple utf-8):\n")

while True:
    if format_fichier in ('txt', 'csv'):
        separator = input(
            'Merci de renseigner le séparateur utilisé pour le fichier à plat:\ntabulation, ",", ";"" ou "|" ?: \n')
        if separator == 'tabulation':
            separator == '\t'
            break
        elif separator in (';', ',', '|'):
            break
        else:
            print("Le séparateur du fichier à plat n'est pas valide ! Merci de réessayer.")
    elif format_fichier == 'xlsx':
        break

while True:
    format_montants = input(
        "Les écritures comptables extraites sont comptabilisées au format:\n 1- Débit - Crédit ? (Tapez 1)\n 2- Montant - Sens (Tapez 2):\n")
    if format_montants == '1':
        df_colonnes = pd.DataFrame(columns=df_colonnes_var_1)
        break
    elif format_montants == '2':
        df_colonnes = pd.DataFrame(columns=df_colonnes_var_2)
        break
    else:
        print("Merci d'entrer un chiffre valide")

os.system('cls')

print(
    "Partie 2: Renseignement des informations relatives aux libellés des champs:\n Merci de bien vouloir entrer EXACTEMENT le nom du champs de l'extraction correspondant au champ du FEC\n")

fields = []

if format_montants == '1':
    fields.extend([
        'JournalCode',
        'JournalLib',
        'EcritureNum',
        'EcritureDate',
        'CompteNum',
        'CompteLib',
        'CompAuxNum',
        'CompAuxLib',
        'PieceRef',
        'PieceDate',
        'EcritureLib',
        'Debit',
        'Credit',
        'EcritureLet',
        'DateLet',
        'ValidDate',
        'Montantdevise',
        'Idevise'
    ])
else:
    fields.extend([
        'JournalCode',
        'JournalLib',
        'EcritureNum',
        'EcritureDate',
        'CompteNum',
        'CompteLib',
        'CompAuxNum',
        'CompAuxLib',
        'PieceRef',
        'PieceDate',
        'EcritureLib',
        'Montant',
        'Sens',
        'EcritureLet',
        'DateLet',
        'ValidDate',
        'Montantdevise',
        'Idevise'
    ])


def getInput(field: str) -> str:
    value = input(f"Merci de renseigner le nom du champ correspondant au champ {field} du FEC: \n")
    if len(value) == 0:
        return field
    return value


mapping = {getInput(field): field for field in fields}

os.system('cls')

# SYSTEME D'OUVERTURE

chemin = input(
    "Merci d'entrer ici un chemin valide vers le fichier d'extraction\nVous pouvez copier un chemin valide par un clic droit tout en maintenant la touche MAJ enfoncée sur le fichier d'extraction. Sélectionnez ensuite 'Copier en tant que chemin d'accès' dans la liste.\nChemin: ")

os.system('cls')

chemin = chemin.replace('\\', '\\\\')
chemin = chemin.replace('\"', '')
if format_fichier == 'xlsx':
    df_init = pd.read_excel(chemin, decimal=',')
elif format_fichier == 'csv':
    df_init = pd.read_csv(chemin, sep=';', decimal=',')
elif format_fichier == 'txt':
    df_init = pd.read_csv(chemin, decimal=',', encoding=encoding_fichier, delimiter=separator)
else:
    print("L'analyse sera erronée puisque vous n'avez pas entré de séparateur valide.\n")

df_init = df_init.rename(columns=mapping)
df = pd.concat([df_colonnes, df_init], ignore_index=True)
df = df.replace('|', '')
df = df.replace('\t', '')

# CALL DEF

if verifier_colonnes_vides(df) == "Les champs obligatoires ne contiennent pas de valeurs vides":
    def_2_ok = verifier_colonnes_vides(df)
    variables_success.append(def_2_ok)
else:
    def_2_erreur = verifier_colonnes_vides(df)
    variables_failure.append(def_2_erreur)

tuple_def_3 = verifier_dates(df, liasse)
if tuple_def_3[0] == "Erreur format: Certaines dates de la colonne EcritureDate ne sont pas au format YYYYMMDD":
    def_3_erreur_1 = tuple_def_3[0]
    variables_failure.append(def_3_erreur_1)
else:
    def_3_ok_1 = tuple_def_3[0]
    variables_success.append(def_3_ok_1)
if tuple_def_3[1] == "Erreur format: Certaines dates de la colonne PieceDate ne sont pas au format YYYYMMDD":
    def_3_erreur_2 = tuple_def_3[1]
    variables_failure.append(def_3_erreur_2)
else:
    def_3_ok_2 = tuple_def_3[1]
    variables_success.append(def_3_ok_2)
if tuple_def_3[2] == "Erreur format: Certaines dates de la colonne ValidDate ne sont pas au format YYYYMMDD":
    def_3_erreur_3 = tuple_def_3[2]
    variables_failure.append(def_3_erreur_3)
else:
    def_3_ok_3 = tuple_def_3[2]
    variables_success.append(def_3_ok_3)
if tuple_def_3[3] == "Erreur de date: Certaines dates de ValidDate sont postérieures à la date limite":
    def_3_erreur_4 = tuple_def_3[3]
    variables_failure.append(def_3_erreur_4)
else:
    def_3_ok_4 = tuple_def_3[3]
    variables_success.append(def_3_ok_4)

if verifier_relation_1_1(
        df) == "Erreur de relation 1-1: Il existe plus d'une valeur dans le champ CompteLib pour au moins une valeur du champ CompteNum":
    def_4_erreur = verifier_relation_1_1(df)
    variables_failure.append(def_4_erreur)
else:
    def_4_ok = verifier_relation_1_1(df)
    variables_success.append(def_4_ok)

if verifier_relation_1_1_ecriture(df) == "Erreur de relation 1-1: Il existe plus d'une valeur dans le champ EcritureLib pour au moins une valeur du champ EcritureNum":
    def_5_erreur = verifier_relation_1_1_ecriture(df)
    variables_failure.append(def_5_erreur)
else:
    def_5_ok = verifier_relation_1_1_ecriture(df)
    variables_success.append(def_5_ok)

if verifier_ecriture_lib_numerique(
        df) == "Erreur: Ma colonne 'EcritureLib' contient au moins une valeur uniquement numérique.":
    def_6_erreur = verifier_ecriture_lib_numerique(df)
    variables_failure.append(def_6_erreur)
else:
    def_6_ok = verifier_ecriture_lib_numerique(df)
    variables_success.append(def_6_ok)

sequence_1 = check_values(df)
if sequence_1[0] == "Nombre d'erreurs dans la séquentialité d'EcritureNum par JournalCode : 0":
    def_7_ok = sequence_1[0]
    variables_success.append(def_7_ok)
else:
    def_7_error = sequence_1[0]
    variables_failure.append(def_7_error)

sequence_2 = verifier_sequentialite_piece_ref(df)
if sequence_2[0] == "Nombre d'erreurs dans la séquentialité d'EcritureNum par JournalCode : 0":
    def_8_ok = sequence_2[0]
    variables_success.append(def_8_ok)
else:
    def_8_error = sequence_2[0]
    variables_failure.append(def_8_error)

if format_montants == '2':
    if verifier_colonne_sens(df) == True:
        (ok, err) = verifier_compensation_montants(df, format_montants)
        if err:
            variables_failure.append(err)
        else:
            variables_success.append(ok)
    else:
        def_9_error = "Compensation des montants:la colonne Sens ne contient pas que des valeurs 'D' et 'C'"
        variables_failure.append(def_9_error)
else:
    (ok, err) = verifier_compensation_montants(df, format_montants)
    if err:
        variables_failure.append(err)
    else:
        variables_success.append(ok)

if format_montants == '2':
    if verifier_colonne_sens(df) == True:
        def_10_output = selectionner_top30_comptes(df, format_montants)
        def_10_ok = "La colonne Sens ne contient que des valeurs 'D' et 'C'"
        variables_success.append(def_10_ok)
    else:
        def_10_error = "Erreur: La colonne Sens ne contient pas que des valeurs 'D' et 'C'"
        variables_failure.append(def_10_error)
        def_10_output = [def_10_error]
else:
    def_10_output = selectionner_top30_comptes(df, format_montants)

if verifier_valeurs_interdites_ecriture_lib(df) == "La colonne EcritureLib ne contient pas de valeur interdite.":
    def_11_ok = verifier_valeurs_interdites_ecriture_lib(df)
    variables_success.append(def_11_ok)
else:
    def_11_error = verifier_valeurs_interdites_ecriture_lib(df)
    variables_failure.append(def_11_error)

if format_montants == '1':
    df['Debit'] = df['Debit'].apply(lambda x: locale.atof(str(x)))
    df['Debit'] = df['Debit'].apply(lambda x: locale.format_string("%.2f", x))
    df['Credit'] = df['Credit'].apply(lambda x: locale.atof(str(x)))
    df['Credit'] = df['Credit'].apply(lambda x: locale.format_string("%.2f", x))
else:
    df['Montant'] = df['Montant'].apply(lambda x: locale.atof(str(x)))
    df['Montant'] = df['Montant'].apply(lambda x: locale.format_string("%.2f", x))
# CREATION RAPPORT
doc = Document()

# Variables avec des données de test
tests_success = variables_success
tests_failure = variables_failure
lines_sequentiality = sequence_1[1]
piece_ref_sequentiality = sequence_2[1]
large_amount_lines = def_10_output

# Partie 1 : Tests réussis
doc.add_heading("Tests réalisés avec succès", level=1).bold = True
for test in tests_success:
    doc.add_paragraph(f"• {test}", style="ListBullet")
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Vert

# Partie 2 : Tests échoués
doc.add_page_break()
doc.add_heading("Tests réalisés révélant une erreur", level=1).bold = True
for test in tests_failure:
    doc.add_paragraph(f"• {test}", style="ListBullet")
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Rouge

# Partie 3 : Annexe
doc.add_page_break()
doc.add_heading("Liste des lignes dont la séquentialité d'EcritureNum par code journal est cassée", level=1).bold = True
for line in lines_sequentiality:
    doc.add_paragraph(f"• {line}", style="ListBullet")

doc.add_heading("Liste des lignes pour lesquelles la séquentialité de PieceRef par EcritureNum est brisée",
                level=1).bold = True
for line in piece_ref_sequentiality:
    doc.add_paragraph(f"• {line}", style="ListBullet")

doc.add_heading("Liste des lignes comprenant les montants les plus importants", level=1).bold = True
for line in large_amount_lines:
    doc.add_paragraph(f"• {line}", style="ListBullet")

chemin_docx = WindowsPath(os.path.dirname(argv[0])).joinpath(f"Rapport {siren}FEC{periode_end}.docx")
# Sauvegarder le document
doc.save(chemin_docx)

os.system('cls')

# GENERATION FEC
chemin_fec = WindowsPath(os.path.dirname(argv[0])).joinpath(f"{siren}FEC{periode_end}.txt")
df = df.sort_values(['ValidDate', 'EcritureNum', 'CompteNum'])
df.to_csv(chemin_fec, sep='|', encoding="utf-8", index=False)
print(
    "Votre extraction au format FEC a été générée dans le même dossier que cet exécutable, il vous revient de la "
    "corriger selon le rapport également généré. N'hésitez pas à compléter cette analyse d'un test TestComptaDemat'.")
